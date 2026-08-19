"""LangGraph 节点实现。

每个节点接收完整 AgentState，返回**局部更新字典**（LangGraph 会按 reducer 合并）。
节点内部只做一件事，路由判断交给 graph.py 的条件边。
"""

from __future__ import annotations

import json
import logging
import re
import sys
from datetime import datetime, timedelta
from pathlib import Path
from typing import Dict, List, Sequence

from langgraph.types import interrupt

from src.agent import prompts as P
from src.agent.llm import chat, chat_json, chat_json_many, chat_many
from src.agent.state import AgentState, Evidence
from src.agent import citation_graph as CG  # 引用网络分析（方向 D'）
from src.agent import quality as QL          # 质量评估仪表盘（方向 F'）
from src.agent.tools import (
    apply_relevance_gate,
    compute_relevance,
    dedup_papers,
    enrich_topn_fulltext,
    hydrate_from_store,
    multi_source_search,
    rank_papers,
    recall_from_store,
    save_to_store,
    year_histogram,
    year_range,
)
from src.cluster.theme_cluster import cluster_papers
from src.config import get_settings
from src.ingest.base import make_paper
from src.ingest.openalex import enrich_citations
from src.report.bibtex import build_bibtex, build_reference_list

logger = logging.getLogger(__name__)

EXTRACT_BATCH_SIZE = 5          # 每次 LLM 调用抽取的论文数
MAX_EXTRACT_PAPERS = 60         # 抽取上限，控制 token 成本
MAX_EVIDENCE_PER_SECTION = 15   # 单小节最多喂给 writer 的证据条数
POOL_HARD_CAP = 120             # 文献池硬上限


def _log(msg: str) -> str:
    stamp = datetime.now().strftime("%H:%M:%S")
    line = f"[{stamp}] {msg}"
    logger.info(msg)
    return line


def _lang_hint() -> str:
    return P.LANG_HINT.get(get_settings().report_language, P.LANG_HINT["zh"])


# ==========================================================================
# 增量更新辅助（方向 B'）
# ==========================================================================
def _parse_since_year(since_date: Optional[str], default_days: int) -> Optional[int]:
    """把 since_date（YYYY-MM-DD）或默认回看窗口换算成年份下限。"""
    if since_date:
        digits = "".join(ch for ch in str(since_date) if ch.isdigit())
        if len(digits) >= 4:
            return int(digits[:4])
    cutoff = datetime.now() - timedelta(days=default_days)
    return cutoff.year


def _load_previous(base_path: str):
    """载入上一版成稿的文献池 / 引用编号 / 小节正文 / 主题簇。

    入参 ``base_path`` 为上一版的 `.md` 成稿（或其同名 `.json` / 目录）；
    从同目录的 ``<stem>_papers.json`` 与 ``<stem>_meta.json`` 还原信息。
    返回 ``(prev_papers, prev_citation_map, prev_sections, prev_clusters)``，
    任一文件缺失则对应项为空，调用方安全降级。
    """
    base = Path(base_path)
    if base.suffix == ".md":
        stem = base.name[: -len("_review.md")] if base.name.endswith("_review.md") else base.stem
        d = base.parent
    else:
        d = base.parent
        stem = base.stem

    papers_file = d / f"{stem}_papers.json"
    meta_file = d / f"{stem}_meta.json"

    prev_papers: List[dict] = []
    prev_citation_map: Dict[str, int] = {}
    if papers_file.exists():
        try:
            for rec in json.loads(papers_file.read_text(encoding="utf-8")):
                prev_papers.append(make_paper(**{k: v for k, v in rec.items()}))
                if rec.get("citation_index") is not None:
                    prev_citation_map[rec["paper_id"]] = int(rec["citation_index"])
        except Exception as exc:
            logger.debug("载入上一版文献失败（跳过）: %s", exc)
            prev_papers, prev_citation_map = [], {}

    prev_sections: Dict[str, str] = {}
    prev_clusters: List[dict] = []
    if meta_file.exists():
        try:
            meta = json.loads(meta_file.read_text(encoding="utf-8"))
            prev_sections = meta.get("sections", {}) or {}
            prev_clusters = meta.get("clusters", []) or []
        except Exception as exc:
            logger.debug("载入上一版 meta 失败（跳过）: %s", exc)

    return prev_papers, prev_citation_map, prev_sections, prev_clusters


# ==========================================================================
# 1. QueryExpander
# ==========================================================================
def query_expander(state: AgentState) -> dict:
    """把研究主题扩展成多组检索式；Critic 打回时只生成补缺口的新检索式。"""
    topic = state["topic"]
    old_queries = state.get("queries") or []
    critic = state.get("critic") or {}
    is_refine = bool(old_queries) and critic.get("verdict") == "need_more"

    if is_refine:
        missing = critic.get("missing_topics") or []
        hinted = critic.get("extra_queries") or []
        user = P.QUERY_EXPANDER_REFINE_USER.format(
            topic=topic,
            old_queries="\n".join(f"- {q}" for q in old_queries),
            missing="\n".join(f"- {m}" for m in missing) or "- （未指明，请自行判断薄弱面）",
        )
        data = chat_json("query_expander", P.QUERY_EXPANDER_SYSTEM, user, default={})
        new_queries = list(data.get("queries") or []) + hinted
    else:
        constraints = (state.get("constraints") or "").strip()
        block = f"额外约束：{constraints}\n\n" if constraints else ""
        user = P.QUERY_EXPANDER_USER.format(topic=topic, constraints_block=block)
        data = chat_json("query_expander", P.QUERY_EXPANDER_SYSTEM, user, default={})
        new_queries = list(data.get("queries") or [])

    # 清洗 + 去重（大小写不敏感）
    seen = {q.lower().strip() for q in old_queries}
    pending: List[str] = []
    for q in new_queries:
        q = re.sub(r"\s+", " ", str(q)).strip().strip('"')
        if q and q.lower() not in seen and 2 <= len(q) <= 120:
            seen.add(q.lower())
            pending.append(q)

    if not pending and not old_queries:  # LLM 失败兜底
        pending = [topic, f"{topic} survey", f"{topic} review"]

    return {
        "queries": old_queries + pending,
        "pending_queries": pending,
        "logs": [_log(f"QueryExpander: {'补充' if is_refine else '生成'} {len(pending)} 条检索式 → {pending}")],
    }


# ==========================================================================
# 2. Retriever
# ==========================================================================
def retriever(state: AgentState) -> dict:
    """并发打多源学术 API 并与已有文献池合并去重。"""
    settings = get_settings()
    round_no = state.get("retrieval_round", 0)
    pending = state.get("pending_queries") or state.get("queries") or []

    # 数量不足触发的重检索：换用全量检索式并放大每条的返回上限
    if not pending:
        pending = state.get("queries") or [state["topic"]]

    # 跨主题复用：先从本地文献池召回与检索式相近的历史文献，再打网络
    seed = recall_from_store(pending, settings)
    if seed:
        logger.info("Retriever: 从本地文献池召回 %d 条历史文献", len(seed))

    # ---- 增量更新（方向 B'）：载入上一版文献池、沿用引用编号 ----
    inc_updates: dict = {}
    if state.get("incremental") and not state.get("previous_loaded") and state.get("base_path"):
        prev_papers, prev_citation_map, prev_sections, prev_clusters = _load_previous(
            state["base_path"]
        )
        if prev_papers:
            seed = (seed or []) + prev_papers
            merged_citation = dict(state.get("citation_map") or {})
            merged_citation.update(prev_citation_map)
            inc_updates = {
                "citation_map": merged_citation,
                "previous_loaded": True,
                "previous_pids": [p["paper_id"] for p in prev_papers],
                "previous_sections": prev_sections,
                "previous_clusters": prev_clusters,
            }
            logger.info("Retriever(增量): 载入上一版 %d 篇文献，沿用历史引用编号", len(prev_papers))

    original_limit = settings.max_results_per_query
    if round_no > 0:
        settings.max_results_per_query = min(100, original_limit * (round_no + 1))

    try:
        fresh = multi_source_search(pending, settings)
    finally:
        settings.max_results_per_query = original_limit

    # ---- 增量更新（方向 B'）：仅保留 since_date 之后的新文献 ----
    # 注意：仅增量模式才过滤日期；常规运行（incremental=False）不应用任何
    # 时间窗口，否则默认回看窗口会把绝大部分历史文献误删为空池。
    if state.get("incremental"):
        since_year = _parse_since_year(state.get("since_date"), settings.incremental_default_days)
        if since_year:
            before = len(fresh)
            fresh = [p for p in fresh if (p.get("year") or 0) >= since_year]
            if before != len(fresh):
                logger.info("Retriever(增量): since_date 过滤后保留 %d/%d 篇（>=%d）",
                            len(fresh), before, since_year)

    merged = dedup_papers(list(state.get("papers") or []) + seed + fresh)

    # 用本地缓存回填引用数/摘要/DOI，并写回本轮文献池（跨运行复用）
    merged = hydrate_from_store(merged)
    save_to_store(merged)

    return {
        **inc_updates,
        "papers": merged,
        "pending_queries": [],
        "retrieval_round": round_no + 1,
        "logs": [
            _log(
                f"Retriever(第 {round_no + 1} 轮): 新检索 {len(fresh)} 条"
                + (f"（本地召回 {len(seed)} 条）" if seed else "")
                + f"，合并去重后文献池 {len(merged)} 篇"
            )
        ],
    }


# ==========================================================================
# 3. Ranker
# ==========================================================================
def ranker(state: AgentState) -> dict:
    """去重后排序 + 相关性闸门 + 补引用数 + 对 Top-N 拉全文。"""
    settings = get_settings()
    papers = list(state.get("papers") or [])
    queries = state.get("queries") or []

    # arXiv 源没有引用数，用 OpenAlex 补齐，让排序信号完整
    if "openalex" in settings.enabled_sources:
        enrich_citations(papers, limit=30)

    # --- 引用网络分析（方向 D'）：PageRank 枢纽度 + 桥接度 ---
    CG.score_centrality(papers)
    graph_available = sum(len(p.get("referenced_works") or []) for p in papers) > 0
    top_hub = sorted(papers, key=lambda p: p.get("hub_score", 0.0), reverse=True)[:5]
    top_bridge = sorted(papers, key=lambda p: p.get("bridge_score", 0.0), reverse=True)[:5]

    # --- 相关性闸门（P0-1）---
    relevance = compute_relevance(papers, state["topic"], queries)
    kept, dropped = apply_relevance_gate(
        papers, relevance, settings.relevance_gate, settings.min_pool_after_gate
    )
    gate_stats = {
        "total": len(papers),
        "kept": len(kept),
        "dropped": len(dropped),
        "threshold": settings.relevance_gate,
        "dropped_titles": dropped[:25],
    }

    dropped_high_hub = [
        {"title": p.get("title", ""), "hub": round(p.get("hub_score", 0.0), 4)}
        for p in dropped if (p.get("hub_score") or 0.0) >= 0.6
    ][:10]

    # 排序用闸门幸存者；relevance 已对齐 kept 顺序，避免重复计算
    score_map = {p["paper_id"]: relevance[i] for i, p in enumerate(papers)}
    kept_relevance = [score_map[p["paper_id"]] for p in kept]
    ranked = rank_papers(kept, state["topic"], queries, settings, relevance=kept_relevance)[
        :POOL_HARD_CAP
    ]
    n_full = enrich_topn_fulltext(ranked, settings)

    citation_analysis = {
        "available": graph_available,
        "top_hub": [
            {"title": p.get("title", ""), "year": p.get("year"),
             "hub": round(p.get("hub_score", 0.0), 4), "paper_id": p["paper_id"]}
            for p in top_hub
        ],
        "top_bridge": [
            {"title": p.get("title", ""), "year": p.get("year"),
             "bridge": round(p.get("bridge_score", 0.0), 4), "paper_id": p["paper_id"]}
            for p in top_bridge
        ],
        "dropped_high_hub": dropped_high_hub,
    }

    top_preview = " | ".join(f"{p['title'][:38]}({p.get('year')})" for p in ranked[:3])
    return {
        "papers": ranked,
        "relevance_gate": gate_stats,
        "citation_analysis": citation_analysis,
        "logs": [
            _log(
                f"Ranker: 相关性闸门移除 {gate_stats['dropped']} 篇偏离主题"
                f"（阈值 {gate_stats['threshold']}），保留 {gate_stats['kept']} 篇入池；"
                f"年份 {year_range(ranked)}，全文解析 {n_full} 篇；Top3: {top_preview}"
            )
        ],
    }


# ==========================================================================
# 4. Extractor
# ==========================================================================
def _paper_block(paper: dict, idx: int) -> str:
    body = paper.get("fulltext") or paper.get("abstract") or ""
    body = body[:5000] if paper.get("fulltext") else body[:1800]
    mode = "全文节选" if paper.get("fulltext") else "摘要"
    return (
        f"### 论文 {idx}\n"
        f"paper_id: {paper['paper_id']}\n"
        f"标题: {paper.get('title')}\n"
        f"年份: {paper.get('year')} | 被引: {paper.get('citation_count', 0)}\n"
        f"{mode}: {body}\n"
    )


def extractor(state: AgentState) -> dict:
    """分批从摘要/全文抽取「方法、结论、数据集、指标」，绑定 paper_id。"""
    papers = state.get("papers") or []
    done_ids = {e["paper_id"] for e in (state.get("evidence") or [])}
    targets = [p for p in papers[:MAX_EXTRACT_PAPERS] if p["paper_id"] not in done_ids]

    if not targets:
        return {"logs": [_log("Extractor: 无新增文献，跳过")]}

    # 把每个 batch 打包成一次独立的 LLM 调用，整体并发执行（P3-4）
    batches: List[List[dict]] = []
    items: List[Dict[str, Any]] = []
    for start in range(0, len(targets), EXTRACT_BATCH_SIZE):
        batch = targets[start : start + EXTRACT_BATCH_SIZE]
        block = "\n".join(_paper_block(p, i + 1) for i, p in enumerate(batch))
        user = P.EXTRACTOR_USER.format(count=len(batch), papers_block=block)
        items.append({"node": "extractor", "system": P.EXTRACTOR_SYSTEM, "user": user})
        batches.append(batch)

    results = chat_json_many(items, default={})

    new_evidence: List[Evidence] = []
    for batch, data in zip(batches, results):
        got = {}
        for item in (data.get("evidence") or []):
            pid = str(item.get("paper_id", "")).strip()
            if pid:
                got[pid] = item

        for paper in batch:
            item = got.get(paper["paper_id"], {})
            claim = (item.get("claim") or "").strip()
            if not claim:  # 抽取失败时用摘要首句兜底，保证每篇都有可用证据
                claim = (paper.get("abstract") or paper.get("title") or "")[:200]
            new_evidence.append(
                {
                    "paper_id": paper["paper_id"],
                    "claim": claim,
                    "method": (item.get("method") or "").strip(),
                    "dataset": (item.get("dataset") or "").strip(),
                    "metric": (item.get("metric") or "").strip(),
                    "section": "",
                }
            )

    return {
        "evidence": new_evidence,
        "logs": [_log(f"Extractor: 抽取 {len(new_evidence)} 篇文献的结构化证据")],
    }


# ==========================================================================
# 5. Clusterer
# ==========================================================================
def clusterer(state: AgentState) -> dict:
    """embedding + KMeans 分主题簇，再由 LLM 给每簇起小节标题，并分配引用编号。"""
    settings = get_settings()
    papers = state.get("papers") or []
    covered = {e["paper_id"] for e in (state.get("evidence") or [])}
    pool = [p for p in papers if p["paper_id"] in covered] or papers

    clusters = cluster_papers(pool, n_clusters=settings.n_clusters)
    if not clusters:
        return {"clusters": [], "citation_map": {}, "logs": [_log("Clusterer: 文献池为空")]}

    # --- LLM 命名 ---
    by_id = {p["paper_id"]: p for p in pool}
    blocks = []
    for c in clusters:
        titles = [by_id[pid]["title"] for pid in c.paper_ids[:8] if pid in by_id]
        blocks.append(
            f"簇 {c.cluster_id}（{c.size} 篇）\n"
            f"关键词: {', '.join(c.keywords[:8]) or '（无）'}\n"
            + "\n".join(f"  - {t[:110]}" for t in titles)
        )
    user = P.CLUSTER_NAMER_USER.format(topic=state["topic"], clusters_block="\n\n".join(blocks))
    data = chat_json(
        "cluster_namer",
        P.CLUSTER_NAMER_SYSTEM.format(lang_hint=_lang_hint()),
        user,
        default={},
    )
    labels = {str(k): str(v) for k, v in (data.get("labels") or {}).items()}

    used_labels = set()
    for c in clusters:
        raw = labels.get(str(c.cluster_id), "").strip()
        if not raw or raw in used_labels:
            raw = f"主题 {c.cluster_id + 1}：{', '.join(c.keywords[:3]) or '其他相关研究'}"
        used_labels.add(raw)
        c.label = raw

    # --- 引用编号：跨轮次稳定（方向 A）---
    # 保留上一轮已分配的编号，仅对新出现的论文追加，避免 Critic 打回重聚类后
    # 用户看到的 [n] 指向完全不同的论文。
    existing_map = dict(state.get("citation_map") or {})
    citation_map: Dict[str, int] = dict(existing_map)
    counter = (max(existing_map.values()) if existing_map else 0) + 1
    for c in clusters:
        ordered = sorted(
            (pid for pid in c.paper_ids if pid in by_id),
            key=lambda pid: by_id[pid].get("score", 0.0),
            reverse=True,
        )
        c.paper_ids = ordered
        for pid in ordered:
            if pid not in citation_map:
                citation_map[pid] = counter
                counter += 1

    return {
        "clusters": [c.to_dict() for c in clusters],
        "citation_map": citation_map,
        "logs": [
            _log(
                "Clusterer: "
                + " / ".join(f"{c.label}({c.size})" for c in clusters)
            )
        ],
    }


# ==========================================================================
# 5.5 IncrementalPlan —— 增量更新规划（方向 B'）
# ==========================================================================
def incremental_plan(state: AgentState) -> dict:
    """增量更新规划：对比本版主题簇与上一版，决定保留/重写哪些小节。

    - 非增量模式：直接放行，不改任何状态；
    - 增量模式：把「本版簇」与「上一版簇」按论文重叠度匹配，重叠 >=50% 的小节
      沿用上一版正文（省 LLM token），其余小节标记重写；同时统计新增论文数，
      写入 `incremental_note` 供成稿渲染。
    """
    if not state.get("incremental"):
        return {}

    prev_clusters = state.get("previous_clusters") or []
    prev_sections = state.get("previous_sections") or {}
    prev_pids = set(state.get("previous_pids") or [])
    current_pids = {p["paper_id"] for p in (state.get("papers") or [])}
    new_pids = current_pids - prev_pids

    kept: List[str] = []
    carry: Dict[str, str] = {}
    for c in state.get("clusters") or []:
        c_pids = set(c.get("paper_ids") or [])
        best, best_ratio = None, 0.0
        for pc in prev_clusters:
            pc_pids = set(pc.get("paper_ids") or [])
            if not c_pids:
                continue
            ratio = len(c_pids & pc_pids) / len(c_pids)
            if ratio > best_ratio:
                best, best_ratio = pc, ratio
        label = c["label"]
        if best is not None and best_ratio >= 0.5 and best.get("label") in prev_sections:
            kept.append(label)
            carry[label] = prev_sections[best["label"]]
        else:
            logger.debug("IncrementalPlan: 小节 %r 将重写（重叠 %.2f）", label, best_ratio)

    note = {
        "new": len(new_pids),
        "rewritten": len(state.get("clusters") or []) - len(kept),
        "kept": len(kept),
        "base": state.get("base_path"),
    }
    return {
        "incremental_keep": kept,
        "sections": carry,
        "incremental_note": note,
        "logs": [
            _log(
                f"IncrementalPlan: 新增 {note['new']} 篇，保留 {note['kept']} 个小节、"
                f"重写 {note['rewritten']} 个小节（沿用历史编号）"
            )
        ],
    }


# ==========================================================================
# 6. SectionWriter
# ==========================================================================
def _evidence_block(
    paper_ids: Sequence[str],
    by_id: Dict[str, dict],
    ev_by_id: Dict[str, Evidence],
    citation_map: Dict[str, int],
) -> str:
    lines = []
    for pid in paper_ids[:MAX_EVIDENCE_PER_SECTION]:
        paper = by_id.get(pid)
        if not paper:
            continue
        num = citation_map.get(pid)
        ev = ev_by_id.get(pid, {})
        parts = [
            f"[{num}] {paper.get('title')} ({paper.get('year')}, 被引 {paper.get('citation_count', 0)})",
            f"    结论: {ev.get('claim') or (paper.get('abstract') or '')[:180]}",
        ]
        if ev.get("method"):
            parts.append(f"    方法: {ev['method']}")
        if ev.get("dataset"):
            parts.append(f"    数据: {ev['dataset']}")
        if ev.get("metric"):
            parts.append(f"    指标: {ev['metric']}")
        lines.append("\n".join(parts))
    return "\n\n".join(lines)


def section_writer(state: AgentState) -> dict:
    """为每个主题簇生成带内联引用的综述段落。"""
    clusters = state.get("clusters") or []
    papers = state.get("papers") or []
    by_id = {p["paper_id"]: p for p in papers}
    ev_by_id = {e["paper_id"]: e for e in (state.get("evidence") or [])}
    citation_map = state.get("citation_map") or {}

    # 增量更新（方向 B'）：沿用上一版正文的小节直接保留，不参与重新生成
    keep_set = set(state.get("incremental_keep") or []) if state.get("incremental") else set()
    sections: Dict[str, str] = dict(state.get("sections") or {})
    specs: List[tuple] = []  # (cluster_label, allowed_set, llm_item)
    for cluster in clusters:
        label = cluster["label"]
        if label in keep_set:
            continue  # 沿用上一版正文
        pids = cluster["paper_ids"]
        allowed = [citation_map[pid] for pid in pids[:MAX_EVIDENCE_PER_SECTION] if pid in citation_map]
        if not allowed:
            continue
        user = P.SECTION_WRITER_USER.format(
            topic=state["topic"],
            section_title=cluster["label"],
            keywords=", ".join(cluster.get("keywords") or []) or "（无）",
            allowed_citations=", ".join(f"[{n}]" for n in allowed),
            evidence_block=_evidence_block(pids, by_id, ev_by_id, citation_map),
        )
        specs.append((
            cluster["label"],
            set(allowed),
            {
                "node": "section_writer",
                "system": P.SECTION_WRITER_SYSTEM.format(lang_hint=_lang_hint()),
                "user": user,
            },
        ))

    # 各小节互相独立，整体并发生成（P3-4）
    texts = chat_many([s[2] for s in specs]) if specs else []
    for (label, allowed_set, _), text in zip(specs, texts):
        text = (text or "").strip()
        if text:
            sections[label] = _strip_invalid_citations(text, allowed_set)

    return {
        "sections": sections,
        "logs": [_log(f"SectionWriter: 生成 {len(sections)} 个小节，共 {sum(len(v) for v in sections.values())} 字")],
    }


# ==========================================================================
# 6.5 GroundClaims —— Claim 级证据锚定（P3-2）
# ==========================================================================
def ground_claims(state: AgentState) -> dict:
    """把每个小节中的核心论断拆成 (text, paper_ids, confidence) 三元组。

    让综述从「段落 + [n] 引用」升级到「逐条论断可审计」：每条 claim 都绑定支撑它的
    具体论文与证据强度，便于读者判断可信度，也方便下游做引用防幻觉二次校验。
    """
    sections = state.get("sections") or {}
    citation_map = state.get("citation_map") or {}
    ev_by_id = {e["paper_id"]: e for e in (state.get("evidence") or [])}
    by_id = {p["paper_id"]: p for p in (state.get("papers") or [])}
    num_to_pid = {n: pid for pid, n in citation_map.items()}

    all_grounded: List[Dict[str, Any]] = []
    specs: List[tuple] = []  # (title, llm_item)
    for title, body in sections.items():
        used_nums = sorted({int(m) for m in re.findall(r"\[(\d{1,3})\]", body)})
        allowed_pids = [num_to_pid[n] for n in used_nums if n in num_to_pid]
        if not allowed_pids:
            continue
        ev_lines = []
        for pid in allowed_pids:
            ev = ev_by_id.get(pid, {})
            num = citation_map.get(pid)
            head = f"[{num}] {by_id.get(pid, {}).get('title', '')}"
            claim = (ev.get("claim") or "").strip()
            ev_lines.append(f"{head}: {claim[:200]}" if claim else head)
        user = P.GROUND_CLAIMS_USER.format(
            section_title=title,
            evidence_block="\n".join(ev_lines) or "（无）",
        )
        specs.append((
            title,
            {
                "node": "ground_claims",
                "system": P.GROUND_CLAIMS_SYSTEM.format(lang_hint=_lang_hint()),
                "user": user,
            },
        ))

    # 各小节互相独立，整体并发执行（P3-4）
    results = chat_json_many([s[1] for s in specs], default={}) if specs else []
    for (title, _), data in zip(specs, results):
        data = data or {}
        claims: List[Dict[str, Any]] = []
        for c in (data.get("claims") or []):
            text = str(c.get("text", "")).strip()
            conf = str(c.get("confidence", "medium")).lower()
            if conf not in ("high", "medium", "low"):
                conf = "medium"
            valid: List[str] = []
            for raw in (c.get("paper_ids") or []):
                pid = str(raw).strip().lstrip("#")
                if pid in citation_map:
                    valid.append(pid)
                elif pid.isdigit() and int(pid) in num_to_pid:
                    valid.append(num_to_pid[int(pid)])
            valid = list(dict.fromkeys(valid))
            if text and valid:
                claims.append({"text": text, "paper_ids": valid, "confidence": conf})
        if claims:
            all_grounded.append({"section": title, "claims": claims})

    total = sum(len(g["claims"]) for g in all_grounded)
    return {
        "grounded_claims": all_grounded,
        "logs": [
            _log(
                f"GroundClaims: 为 {len(all_grounded)} 个小节标注 {total} 条 claim 级证据锚定"
            )
        ],
    }


def faithfulness(state: AgentState) -> dict:
    """引用-论断一致性校验（LLM-as-Judge，方向 B）。

    把 `grounded_claims` 中每条 claim 与其支撑论文证据交给 LLM 判定是否真被支持，
    输出整体一致性得分与疑似无支撑的论断列表，写入 `faithfulness`，
    供成稿附录 A.7 呈现与告警，形成「生成 + 自检」闭环。
    """
    settings = get_settings()
    grounded = state.get("grounded_claims") or []
    evidence = state.get("evidence") or []
    by_id = {e["paper_id"]: e for e in evidence}

    if not settings.enable_faithfulness:
        return {"faithfulness": {"skipped": True},
                "logs": [_log("Faithfulness: 已关闭（ENABLE_FAITHFULNESS=false），跳过")]}
    if not grounded:
        return {"faithfulness": {"score": 1.0, "checked": 0, "flagged": [], "skipped": True},
                "logs": [_log("Faithfulness: 无 claim 可校验，跳过")]}

    specs: List[tuple] = []
    for g in grounded:
        section = g.get("section", "")
        for c in g.get("claims", []):
            ev_lines = []
            for pid in c.get("paper_ids", []):
                ev = by_id.get(pid, {})
                txt = (ev.get("claim") or "").strip()
                if txt:
                    ev_lines.append(f"- [{pid}] {txt[:300]}")
            user = P.FAITHFULNESS_USER.format(
                claim=c.get("text", ""),
                evidence_block="\n".join(ev_lines) or "（无支撑证据）",
            )
            specs.append((
                section, c.get("text", ""),
                {
                    "node": "faithfulness",
                    "system": P.FAITHFULNESS_SYSTEM.format(lang_hint=_lang_hint()),
                    "user": user,
                },
            ))

    if not specs:
        return {"faithfulness": {"score": 1.0, "checked": 0, "flagged": [], "skipped": True},
                "logs": [_log("Faithfulness: 无 claim 可校验，跳过")]}

    results = chat_json_many([s[2] for s in specs], default={}) if specs else []
    flagged: List[Dict[str, Any]] = []
    supported = 0
    for spec, data in zip(specs, results):
        data = data or {}
        verdict = str(data.get("verdict", "supported")).lower()
        reason = str(data.get("reason", ""))
        if verdict.startswith("unsupp"):
            flagged.append({"section": spec[0], "text": spec[1], "reason": reason})
        else:
            supported += 1
    checked = len(specs)
    score = round(supported / checked, 3) if checked else 1.0
    return {
        "faithfulness": {"score": score, "checked": checked, "flagged": flagged},
        "logs": [_log(
            f"Faithfulness: 校验 {checked} 条 claim，{len(flagged)} 条疑似无充分支撑"
            f"（一致性得分 {score}）"
        )],
    }


def _strip_invalid_citations(text: str, allowed: set) -> str:
    """删除模型编造的、不在候选列表中的引用编号，保证引用可追溯。"""

    def repl(m: re.Match) -> str:
        num = int(m.group(1))
        return m.group(0) if num in allowed else ""

    return re.sub(r"\[(\d{1,3})\]", repl, text)


# ==========================================================================
# 7. Critic
# ==========================================================================
def critic(state: AgentState) -> dict:
    """评审覆盖度与矛盾处理，决定补文献还是放行。"""
    settings = get_settings()
    sections = state.get("sections") or {}
    papers = state.get("papers") or []
    round_no = state.get("critic_round", 0)

    if not sections:
        return {
            "critic": {"verdict": "pass", "coverage_score": 0, "comments": "无小节内容，直接放行"},
            "critic_round": round_no + 1,
            "logs": [_log("Critic: 无小节内容，跳过评审")],
        }

    block = "\n\n".join(f"### {title}\n{body}" for title, body in sections.items())
    user = P.CRITIC_USER.format(
        topic=state["topic"],
        paper_count=len(papers),
        year_range=year_range(papers),
        critic_round=round_no + 1,
        max_rounds=settings.max_critic_rounds,
        sections_block=block[:16000],
    )
    report = chat_json("critic", P.CRITIC_SYSTEM, user, default={}) or {}

    verdict = str(report.get("verdict", "pass")).lower()
    report["verdict"] = "need_more" if verdict.startswith("need") else "pass"
    try:
        report["coverage_score"] = int(report.get("coverage_score", 7))
    except (TypeError, ValueError):
        report["coverage_score"] = 7

    return {
        "critic": report,
        "critic_round": round_no + 1,
        "logs": [
            _log(
                f"Critic(第 {round_no + 1} 轮): {report['verdict']}，"
                f"覆盖度 {report['coverage_score']}/10，"
                f"缺口 {report.get('missing_topics') or '无'}"
            )
        ],
    }


# ==========================================================================
# 8. GapAnalyzer
# ==========================================================================
def gap_analyzer(state: AgentState) -> dict:
    """识别研究空白与技术演进趋势。"""
    sections = state.get("sections") or {}
    papers = state.get("papers") or []
    clusters = state.get("clusters") or []

    hist = year_histogram(papers)
    hist_str = ", ".join(f"{y}:{n}" for y, n in list(hist.items())[-12:]) or "未知"
    cluster_summary = "; ".join(f"{c['label']}({c['size']}篇)" for c in clusters) or "无"
    block = "\n\n".join(f"### {t}\n{b}" for t, b in sections.items())

    user = P.GAP_ANALYZER_USER.format(
        topic=state["topic"],
        year_hist=hist_str,
        cluster_summary=cluster_summary,
        critic_comments=(state.get("critic") or {}).get("comments", "（无）"),
        sections_block=block[:16000],
    )
    data = chat_json(
        "gap_analyzer",
        P.GAP_ANALYZER_SYSTEM.format(lang_hint=_lang_hint()),
        user,
        default={},
    ) or {}

    gaps = [str(g).strip() for g in (data.get("gaps") or []) if str(g).strip()]
    trends = [str(t).strip() for t in (data.get("trends") or []) if str(t).strip()]

    # 共引分析补充潜在空白（方向 D'）：把共享大量参考文献的论文聚成子领域，
    # 未被现有主题簇覆盖者标为研究空白候选，与 LLM 识别的 gaps 合并去重。
    try:
        cocite_gaps = CG.cocitation_gaps(papers, clusters)
    except Exception as exc:
        logger.debug("共引空白分析失败（跳过）: %s", exc)
        cocite_gaps = []
    for g in cocite_gaps:
        if g and g not in gaps:
            gaps.append(g)

    return {
        "gaps": gaps,
        "trends": trends,
        "logs": [_log(f"GapAnalyzer: {len(gaps)} 条研究空白（含共引 {len(cocite_gaps)} 条），{len(trends)} 条趋势")],
    }


# ==========================================================================
# 9. Synthesizer
# ==========================================================================
def _slugify(text: str) -> str:
    s = re.sub(r"[^\w\u4e00-\u9fff]+", "_", text).strip("_")
    return s[:60] or "review"


def synthesizer(state: AgentState) -> dict:
    """汇总成稿：摘要 + 引言 + 各主题小节 + 空白趋势 + 结论 + 参考文献 + BibTeX。"""
    settings = get_settings()
    settings.ensure_dirs()

    topic = state["topic"]
    sections = state.get("sections") or {}
    papers = state.get("papers") or []
    citation_map = state.get("citation_map") or {}
    gaps = state.get("gaps") or []
    trends = state.get("trends") or []
    lang_hint = _lang_hint()

    cited_papers = [p for p in papers if p["paper_id"] in citation_map]
    allowed = sorted(citation_map.values())
    allowed_str = ", ".join(f"[{n}]" for n in allowed[:80])
    section_titles = "\n".join(f"- {t}" for t in sections) or "- （无）"

    # --- 摘要 + 引言 ---
    digest = "\n\n".join(f"### {t}\n{b[:400]}" for t, b in sections.items())
    head_raw = chat(
        "synthesizer",
        P.SYNTH_ABSTRACT_SYSTEM.format(lang_hint=lang_hint),
        P.SYNTH_ABSTRACT_USER.format(
            topic=topic,
            paper_count=len(cited_papers),
            year_range=year_range(cited_papers),
            section_titles=section_titles,
            allowed_citations=allowed_str,
            sections_digest=digest[:12000],
        ),
    )
    abstract, introduction = _split_abstract_intro(head_raw)

    # --- 结论 ---
    conclusion = chat(
        "synthesizer",
        P.SYNTH_CONCLUSION_SYSTEM.format(lang_hint=lang_hint),
        P.SYNTH_CONCLUSION_USER.format(
            topic=topic,
            section_titles=section_titles,
            gaps_block="\n".join(f"- {g}" for g in gaps) or "- （无）",
            trends_block="\n".join(f"- {t}" for t in trends) or "- （无）",
            allowed_citations=allowed_str,
        ),
    ).strip()

    allowed_set = set(allowed)
    abstract = _strip_invalid_citations(abstract, allowed_set)
    introduction = _strip_invalid_citations(introduction, allowed_set)
    conclusion = _strip_invalid_citations(conclusion, allowed_set)

    # --- 组装 Markdown ---
    bibtex, cite_keys = build_bibtex(cited_papers)
    refs = build_reference_list(cited_papers, citation_map, cite_keys)
    report = _assemble_markdown(
        state, abstract, introduction, sections, gaps, trends, conclusion, refs
    )

    # --- 落盘 ---
    slug = _slugify(topic)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    out = settings.output_dir
    paths = {
        "report": out / f"{slug}_{stamp}_review.md",
        "bibtex": out / f"{slug}_{stamp}_references.bib",
        "papers": out / f"{slug}_{stamp}_papers.json",
    }
    paths["report"].write_text(report, encoding="utf-8")
    paths["bibtex"].write_text(bibtex, encoding="utf-8")
    paths["papers"].write_text(
        json.dumps(
            [
                {
                    **{k: v for k, v in p.items() if k != "fulltext"},
                    "has_fulltext": bool(p.get("has_fulltext")),
                    "fulltext_chars": int(p.get("fulltext_chars") or 0),
                    "citation_index": citation_map.get(p["paper_id"]),
                }
                for p in papers
            ],
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )

    # --- 增量更新侧车（方向 B'）：保存本版小节正文与主题簇，供下一版沿用编号 + 保留旧小节 ---
    meta_path = out / f"{slug}_{stamp}_meta.json"
    meta_path.write_text(
        json.dumps(
            {
                "sections": {k: v for k, v in (state.get("sections") or {}).items()},
                "clusters": state.get("clusters") or [],
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )
    paths["meta"] = str(meta_path)

    # --- 引用网络可视化数据（方向 E'）：序列化供 Web UI 渲染交互式网络图 ---
    clusters_for_graph = state.get("clusters") or []
    gap_list = state.get("gaps") or []
    graph_data = CG.export_graph(papers, clusters_for_graph, gap_list)
    graph_path = out / f"{slug}_{stamp}_citation_graph.json"
    graph_path.write_text(
        json.dumps(graph_data, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    paths["citation_graph"] = str(graph_path)

    # --- 质量评估仪表盘（方向 F'）：聚合 B/D'/E'/P3-2 质量信号，落盘供 Web UI 渲染 ---
    quality = QL.compute_quality_report(
        papers=papers,
        clusters=state.get("clusters") or [],
        sections=sections,
        faithfulness=state.get("faithfulness") or {},
        citation_analysis=state.get("citation_analysis") or {},
        citation_graph=graph_data,
        grounded_claims=state.get("grounded_claims") or [],
        gaps=gap_list,
    )
    quality_path = out / f"{slug}_{stamp}_quality_report.json"
    quality_path.write_text(
        json.dumps(quality, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    paths["quality_report"] = str(quality_path)

    # --- 多格式输出（方向 C）：在 Markdown 之外按需再产出 LaTeX / docx ---
    fmt = settings.output_format
    if fmt in ("latex", "docx"):
        bib_basename = f"{slug}_{stamp}_references"
        if fmt == "latex":
            tex = _assemble_latex(
                state, abstract, introduction, sections, gaps, trends, conclusion,
                citation_map, cite_keys, bib_basename,
            )
            tex_path = out / f"{slug}_{stamp}_review.tex"
            tex_path.write_text(tex, encoding="utf-8")
            paths["report_latex"] = str(tex_path)
        if fmt == "docx":
            docx_path = out / f"{slug}_{stamp}_review.docx"
            _write_docx(docx_path, state, abstract, introduction, sections, gaps, trends, conclusion, citation_map)
            paths["report_docx"] = str(docx_path)

    return {
        "report": report,
        "bibtex": bibtex,
        "citation_graph": graph_data,
        "quality_report": quality,
        "artifacts": {k: str(v) for k, v in paths.items()},
        "logs": [
            _log(
                f"Synthesizer: 成稿 {len(report)} 字，引用 {len(cited_papers)} 篇 → {paths['report'].name}"
            )
        ],
    }


def _split_abstract_intro(raw: str) -> tuple[str, str]:
    """从「## 摘要 / ## 引言」格式的输出中拆分两段。"""
    raw = (raw or "").strip()
    m = re.split(r"\n#+\s*(?:引言|Introduction)\s*\n", raw, maxsplit=1, flags=re.IGNORECASE)
    if len(m) == 2:
        abstract = re.sub(r"^#+\s*(?:摘要|Abstract)\s*\n", "", m[0].strip(), flags=re.IGNORECASE)
        return abstract.strip(), m[1].strip()
    return raw, ""


def _assemble_markdown(
    state: AgentState,
    abstract: str,
    introduction: str,
    sections: Dict[str, str],
    gaps: List[str],
    trends: List[str],
    conclusion: str,
    references: str,
) -> str:
    settings = get_settings()
    papers = state.get("papers") or []
    citation_map = state.get("citation_map") or {}
    critic_report = state.get("critic") or {}

    parts: List[str] = []
    cited_papers_all = [p for p in papers if p["paper_id"] in citation_map]
    n_full = sum(1 for p in cited_papers_all if p.get("has_fulltext"))
    if n_full:
        fulltext_note = f"已获取 {n_full} 篇 OA 全文（其余基于摘要）"
    else:
        fulltext_note = "基于摘要撰写（未获取到可用 OA 全文）"
    parts.append(f"# {state['topic']} —— 文献综述\n")

    # 增量更新说明（方向 B'）：在成稿头渲染新增/重写/保留统计
    inc = state.get("incremental_note") or {}
    inc_line = ""
    if state.get("incremental") and inc:
        inc_line = (
            f"> 增量更新：新增 {inc.get('new', 0)} 篇，重写 {inc.get('rewritten', 0)} 个小节，"
            f"保留 {inc.get('kept', 0)} 个小节（沿用历史引用编号）。\n"
        )
    parts.append(
        "> 本文由 **lit-review-agent** 自动生成。\n>\n"
        f"> 生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')} | "
        f"检索源：{', '.join(settings.enabled_sources)} | "
        f"候选文献 {len(papers)} 篇，引用 {len(citation_map)} 篇 | "
        f"年份跨度 {year_range(cited_papers_all)}\n>\n"
        "> 所有论断均带内联引用 `[n]`，编号对应文末参考文献表；"
        f"{fulltext_note}。\n"
        + inc_line
    )

    parts.append("## 摘要\n\n" + (abstract or "（生成失败）"))

    # 章节号连续自增，缺失引言时不留空号
    counter = iter(range(1, 100))
    if introduction:
        parts.append(f"## {next(counter)}. 引言\n\n{introduction}")

    for title, body in sections.items():
        parts.append(f"## {next(counter)}. {title}\n\n{body}")

    if trends or gaps:
        parts.append(f"## {next(counter)}. 研究空白与演进趋势")
        if trends:
            parts.append("### 技术演进趋势\n\n" + "\n".join(f"{i}. {t}" for i, t in enumerate(trends, 1)))
        if gaps:
            parts.append("### 研究空白\n\n" + "\n".join(f"{i}. {g}" for i, g in enumerate(gaps, 1)))

    if conclusion:
        parts.append(f"## {next(counter)}. 结论与展望\n\n{conclusion}")

    parts.append(f"## {next(counter)}. 参考文献\n\n{references or '（无）'}")

    # 附录：让整个生成过程可复核
    appendix = [
        f"## 附录 A：生成过程\n",
        "### A.1 检索式\n",
        "\n".join(f"{i}. `{q}`" for i, q in enumerate(state.get("queries") or [], 1)) or "（无）",
        "\n### A.2 主题簇\n",
        "\n".join(
            f"- **{c['label']}**（{c['size']} 篇）关键词：{', '.join(c.get('keywords') or []) or '—'}"
            for c in (state.get("clusters") or [])
        ) or "（无）",
        "\n### A.3 评审结论\n",
        f"- 判定：`{critic_report.get('verdict', 'n/a')}`，覆盖度 {critic_report.get('coverage_score', 'n/a')}/10",
        f"- 评语：{critic_report.get('comments', '（无）')}",
    ]
    if critic_report.get("contradictions"):
        appendix.append("- 已识别的结论冲突：\n" + "\n".join(
            f"  - {c}" for c in critic_report["contradictions"]))

    # A.5 相关性闸门（P0-1）：让检索质量可审计
    gate = state.get("relevance_gate") or {}
    if gate:
        appendix.append("\n### A.5 相关性闸门\n")
        appendix.append(
            f"- 阈值 `{gate.get('threshold')}`：候选 {gate.get('total')} 篇 → "
            f"移除 **{gate.get('dropped')}** 篇偏离主题，保留 {gate.get('kept')} 篇入池"
        )
        dropped_titles = gate.get("dropped_titles") or []
        if dropped_titles:
            appendix.append("移除的论文（主题相关性不足）：\n" + "\n".join(
                f"  - {t}" for t in dropped_titles[:15]) or "（无）")

    appendix.append("\n### A.4 执行轨迹\n")
    appendix.append("```\n" + "\n".join(state.get("logs") or []) + "\n```")

    # A.6 Claim 级证据锚定（P3-2）：让每条核心论断可追溯、可评强度
    grounded = state.get("grounded_claims") or []
    if grounded:
        appendix.append("\n### A.6 Claim 级证据锚定\n")
        badge = {"high": "● 强", "medium": "◐ 中", "low": "○ 弱"}
        lines = []
        for g in grounded:
            lines.append(f"**{g.get('section', '')}**")
            for c in g.get("claims", []):
                nums = " ".join(
                    f"[{state.get('citation_map', {}).get(pid)}]"
                    for pid in c.get("paper_ids", [])
                    if pid in state.get("citation_map", {})
                )
                conf = c.get("confidence", "medium")
                lines.append(f"- {badge.get(conf, '◐ 中')} {c.get('text', '')} {nums}")
        appendix.append("\n".join(lines) or "（无）")

    # A.7 引用-论断一致性校验（faithfulness，方向 B）：让自检闭环可复核
    faith = state.get("faithfulness") or {}
    if faith and not faith.get("skipped"):
        appendix.append("\n### A.7 引用-论断一致性校验（faithfulness）\n")
        appendix.append(
            f"- 校验论断 {faith.get('checked', 0)} 条，一致性得分 "
            f"**{faith.get('score', 1.0)}**（1.0 = 全部有充分支撑）"
        )
        flagged = faith.get("flagged") or []
        if flagged:
            appendix.append("- 以下论断支撑证据不足，建议人工复核：")
            for f in flagged[:15]:
                appendix.append(
                    f"  - [{f.get('section', '')}] {f.get('text', '')} — {f.get('reason', '')}"
                )
        else:
            appendix.append("- 未发现明显无支撑论断。")

    # A.8 引用网络分析（方向 D'）：让关键文献识别与误删告警可复核
    analysis = state.get("citation_analysis") or {}
    if analysis.get("available") and (analysis.get("top_hub") or analysis.get("top_bridge")):
        appendix.append("\n### A.8 引用网络分析\n")
        appendix.append(
            "- 基于论文间引用关系计算枢纽度（PageRank）与桥接度（betweenness），"
            "识别必引候选与跨子领域枢纽，缓解关键流派漏检。"
        )
        top_hub = analysis.get("top_hub") or []
        if top_hub:
            appendix.append("- 枢纽度 Top（必引候选）：")
            for h in top_hub:
                appendix.append(
                    f"  - {h.get('title', '')} ({h.get('year') or '?'}) — 枢纽度 {h.get('hub')}"
                )
        top_bridge = analysis.get("top_bridge") or []
        if top_bridge:
            appendix.append("- 桥接度 Top（跨子领域枢纽）：")
            for b in top_bridge:
                appendix.append(
                    f"  - {b.get('title', '')} ({b.get('year') or '?'}) — 桥接度 {b.get('bridge')}"
                )
        dropped = analysis.get("dropped_high_hub") or []
        if dropped:
            appendix.append(
                "- 告警：以下被相关性闸门剔除的论文枢纽度较高，建议人工复核是否误删："
            )
            for d in dropped:
                appendix.append(f"  - {d.get('title', '')} — 枢纽度 {d.get('hub')}")

    parts.append("\n".join(appendix))

    return "\n\n".join(parts) + "\n"


def _assemble_latex(
    state: AgentState,
    abstract: str,
    introduction: str,
    sections: Dict[str, str],
    gaps: List[str],
    trends: List[str],
    conclusion: str,
    citation_map: Dict[str, int],
    cite_keys: Dict[str, str],
    bib_basename: str,
) -> str:
    """把结构化成稿字段组装为可编译的 LaTeX 文档（方向 C）。

    与 `_assemble_markdown` 平行：内联 `[n]` 转为 ``\\cite{key}``（key 取自 BibTeX 映射），
    配合同目录已写出的 ``.bib`` 文件即可直接 ``pdflatex`` 编译。
    """
    num_to_pid = {n: pid for pid, n in citation_map.items()}

    def cite(num: int) -> str:
        pid = num_to_pid.get(num)
        key = cite_keys.get(pid) if pid else None
        return f"\\cite{{{key}}}" if key else f"[noref:{num}]"

    def esc(t: str) -> str:
        t = str(t)
        for ch, rep in (
            ("\\", "\\textbackslash{}"), ("&", "\\&"), ("%", "\\%"), ("$", "\\$"),
            ("#", "\\#"), ("_", "\\_"), ("{", "\\{"), ("}", "\\}"),
            ("~", "\\textasciitilde{}"), ("^", "\\textasciicircum{}"),
        ):
            t = t.replace(ch, rep)
        return t

    def inline(t: str) -> str:
        # 先转义特殊字符，再做引用替换（避免 \\cite 的反斜杠被二次转义）
        t = esc(t)
        return re.sub(r"\[(\d{1,3})\]", lambda m: cite(int(m.group(1))), t)

    def paras(text: str) -> str:
        return "\n\n".join(inline(p.strip()) for p in str(text).split("\n\n") if p.strip())

    L: List[str] = [
        "\\documentclass[11pt]{article}",
        "\\usepackage[utf8]{inputenc}",
        "\\usepackage[T1]{fontenc}",
        "\\usepackage[colorlinks=true]{hyperref}",
        "\\usepackage{cite}",
        f"\\title{{{esc(state['topic'])} \\\\ \\large Literature Review (generated by lit-review-agent)}}",
        "\\author{lit-review-agent}",
        "\\begin{document}",
        "\\maketitle",
        "\\begin{abstract}",
        paras(abstract or ""),
        "\\end{abstract}",
    ]
    if introduction:
        L.append("\\section{Introduction}")
        L.append(paras(introduction))
    for title, text in sections.items():
        L.append(f"\\section{{{esc(title)}}}")
        L.append(paras(text))
    if trends or gaps:
        L.append("\\section{Research Gaps and Trends}")
        if trends:
            L.append("\\subsection{Technology Trends}")
            L.append("\\begin{enumerate}")
            for t in trends:
                L.append(f"\\item {inline(t)}")
            L.append("\\end{enumerate}")
        if gaps:
            L.append("\\subsection{Research Gaps}")
            L.append("\\begin{enumerate}")
            for g in gaps:
                L.append(f"\\item {inline(g)}")
            L.append("\\end{enumerate}")
    if conclusion:
        L.append("\\section{Conclusion}")
        L.append(paras(conclusion))
    L.append("\\bibliographystyle{plain}")
    L.append(f"\\bibliography{{{bib_basename}}}")
    L.append("\\end{document}")
    return "\n".join(L) + "\n"


def _write_docx(
    path,
    state: AgentState,
    abstract: str,
    introduction: str,
    sections: Dict[str, str],
    gaps: List[str],
    trends: List[str],
    conclusion: str,
    citation_map: Dict[str, int],
) -> None:
    """把结构化成稿字段写出为 .docx（方向 C，依赖 python-docx）。"""
    try:
        import docx  # python-docx
    except ImportError:  # pragma: no cover - 取决于是否安装 [docx] extra
        raise RuntimeError(
            "docx 输出需要 python-docx，请先运行 `pip install -e \".[docx]\"` 后重试。"
        )
    d = docx.Document()
    d.add_heading(state["topic"] + " — 文献综述", level=0)
    if abstract:
        d.add_paragraph(abstract)
    idx = 1
    if introduction:
        d.add_heading("1. 引言", level=1)
        d.add_paragraph(introduction)
    for title, text in sections.items():
        d.add_heading(f"{idx}. {title}", level=1)
        for p in str(text).split("\n\n"):
            if p.strip():
                d.add_paragraph(p.strip())
        idx += 1
    if trends or gaps:
        d.add_heading(f"{idx}. 研究空白与演进趋势", level=1)
        if trends:
            d.add_heading("技术演进趋势", level=2)
            for t in trends:
                d.add_paragraph(str(t), style="List Number")
        if gaps:
            d.add_heading("研究空白", level=2)
            for g in gaps:
                d.add_paragraph(str(g), style="List Number")
        idx += 1
    if conclusion:
        d.add_heading(f"{idx}. 结论与展望", level=1)
        d.add_paragraph(conclusion)
    d.save(str(path))


# ==========================================================================
# 10. Human（可选人工审核挂起点）
# ==========================================================================
def human_review(state: AgentState) -> dict:
    """人工审核节点（Human-in-the-loop）。

    图运行到此处时调用 `interrupt()` 挂起：把草稿路径交给调用方，等待其用
    `Command(resume=feedback)` 续跑。

    - feedback 为空或近似 'approve' 视为通过：清除 `human_feedback` 并定稿；
    - 否则保留意见到 `human_feedback`，由 `route_after_human` 路由到
      `parse_human_feedback` → `rewrite_sections` 做**针对性重生成**（方向 A），
      而非整篇重来。重生成后再回到本节点等待下一轮审核（受 `max_human_rounds` 上限约束）。
    """
    artifacts = state.get("artifacts") or {}
    report_path = artifacts.get("report")
    payload = {
        "kind": "human_review",
        "message": "综述草稿已生成，请审核后定稿。回复 'approve' 通过，或给出修改意见。",
        "report_path": report_path,
    }
    decision = interrupt(payload)
    feedback = (decision or "").strip()
    if feedback and feedback.lower() not in ("approve", "ok", "yes", "通过", "y", "true"):
        return {
            "human_feedback": feedback,
            "logs": [_log(f"Human: 收到审核意见 → {feedback[:200]}")],
        }
    # 通过：清除可能的旧意见，确保 route_after_human 不会误判为需要改写
    return {"human_feedback": "", "logs": [_log("Human: 审核通过，定稿")]}


def skip_human_review(state: AgentState) -> dict:
    """未启用 --human 时的占位节点：直接放行，不挂起（避免无谓中断）。"""
    return {"logs": [_log("Human: 未启用人工审核，跳过")]}


# ==========================================================================
# 10.5 ParseHumanFeedback —— 把人工意见解析成 targeted 改写动作（方向 A）
# ==========================================================================
def parse_human_feedback(state: AgentState) -> dict:
    """把 `human_feedback` 的自由文本解析为针对现有小节的 rewrite / add 动作列表。

    解析结果写入 `rewrite_targets`，供 `rewrite_sections` 只重跑受影响的小节。
    """
    feedback = (state.get("human_feedback") or "").strip()
    sections = state.get("sections") or {}
    titles = "\n".join(f"- {t}" for t in sections) or "- （无）"
    user = P.PARSE_HUMAN_FEEDBACK_USER.format(
        sections=titles,
        feedback=feedback or "（未提供具体意见）",
    )
    data = chat_json(
        "parse_human_feedback",
        P.PARSE_HUMAN_FEEDBACK_SYSTEM,
        user,
        default={},
    ) or {}

    targets: List[Dict[str, Any]] = []
    for t in (data.get("targets") or []):
        action = str(t.get("action", "rewrite")).lower()
        if action not in ("rewrite", "add"):
            action = "rewrite"
        title = str(t.get("section", "")).strip()
        instr = str(t.get("instruction", "")).strip()
        if title or instr:
            targets.append({"action": action, "section": title, "instruction": instr})

    summary = [t.get("section") or "(新增)" for t in targets] or ["（无）"]
    return {
        "rewrite_targets": targets,
        "logs": [_log(f"ParseHumanFeedback: 解析出 {len(targets)} 条修改动作 → {summary}")],
    }


# ==========================================================================
# 10.6 RewriteSections —— 仅重跑受影响小节（方向 A，闭环重生成核心）
# ==========================================================================
def _best_cluster_match(query: str, clusters: List[Dict[str, Any]]):
    """按关键词重叠为改写动作匹配最相关的主题簇（用于 add 或标题模糊匹配）。"""
    if not query or not clusters:
        return None
    q = set(re.findall(r"[A-Za-z\u4e00-\u9fff]+", (query or "").lower()))
    if not q:
        return None
    best, best_score = None, 0
    for c in clusters:
        text = (c.get("label", "") + " " + " ".join(c.get("keywords") or [])).lower()
        words = set(re.findall(r"[A-Za-z\u4e00-\u9fff]+", text))
        score = len(q & words)
        if score > best_score:
            best, best_score = c, score
    return best


def _generate_section(
    topic: str,
    title: str,
    cluster: Dict[str, Any],
    instruction: str,
    by_id: Dict[str, dict],
    ev_by_id: Dict[str, Evidence],
    citation_map: Dict[str, int],
) -> str:
    """基于某个主题簇的证据，按人工意见重写一篇小节（复用 section_writer 提示词）。"""
    pids = cluster["paper_ids"]
    allowed = [citation_map[pid] for pid in pids[:MAX_EVIDENCE_PER_SECTION] if pid in citation_map]
    if not allowed:
        return ""
    user = P.SECTION_WRITER_USER.format(
        topic=topic,
        section_title=title,
        keywords=", ".join(cluster.get("keywords") or []) or "（无）",
        allowed_citations=", ".join(f"[{n}]" for n in allowed),
        evidence_block=_evidence_block(pids, by_id, ev_by_id, citation_map),
    )
    if instruction:
        user += f"\n\n【人工修改意见】请特别回应以下意见：{instruction}"
    text = chat(
        "section_writer",
        P.SECTION_WRITER_SYSTEM.format(lang_hint=_lang_hint()),
        user,
    ).strip()
    return _strip_invalid_citations(text, set(allowed)) if text else ""


def rewrite_sections(state: AgentState) -> dict:
    """按 `rewrite_targets` 只重跑受影响的小节，并回写 `sections`。

    改写后清掉 `human_feedback`，避免下一轮 approve 被误判为仍需改写；
    同时递增 `human_round` 供 `route_after_human` 做轮次上限控制。
    """
    targets = state.get("rewrite_targets") or []
    sections: Dict[str, str] = dict(state.get("sections") or {})
    clusters = state.get("clusters") or []
    papers = state.get("papers") or []
    by_id = {p["paper_id"]: p for p in papers}
    ev_by_id = {e["paper_id"]: e for e in (state.get("evidence") or [])}
    citation_map = state.get("citation_map") or {}
    cluster_by_label = {c["label"]: c for c in clusters}
    human_round = int(state.get("human_round", 0)) + 1

    for tgt in targets:
        action = (tgt.get("action") or "rewrite").lower()
        title = (tgt.get("section") or "").strip()
        instr = (tgt.get("instruction") or "").strip()
        cluster = cluster_by_label.get(title) or _best_cluster_match(title or instr, clusters)

        if action == "add" or cluster is None:
            cluster = cluster or _best_cluster_match(instr, clusters) or (clusters[0] if clusters else None)
            if cluster is None:
                continue  # 无任何簇可生成，跳过该动作
            new_title = title or cluster["label"]
            sections[new_title] = _generate_section(
                state["topic"], new_title, cluster, instr, by_id, ev_by_id, citation_map
            )
        else:
            sections[title] = _generate_section(
                state["topic"], title, cluster, instr, by_id, ev_by_id, citation_map
            )

    return {
        "sections": sections,
        "human_round": human_round,
        "rewrite_targets": [],
        "human_feedback": "",  # 清除，避免下一轮 approve 误判
        "logs": [
            _log(
                f"RewriteSections: 按人工意见改写 {len(targets)} 个小节（第 {human_round} 轮，"
                f"当前小节数 {len(sections)}）"
            )
        ],
    }
